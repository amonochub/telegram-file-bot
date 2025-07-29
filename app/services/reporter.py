import shutil

import docx
import fitz
import structlog

from app.services.comparer import compare_tokens
from app.services.extractor import extract_pairs
from app.services.tokeniser import extract_tokens

log = structlog.get_logger(__name__)


def validate_doc(path: str):
    try:
        pairs = extract_pairs(path)
        all_miss = []
        for i, (l, r) in enumerate(pairs, 1):
            miss = compare_tokens(extract_tokens(l), extract_tokens(r))
            if miss:
                all_miss.append((i, l, r))
        patched = highlight_diffs(path, all_miss)
        log.info("doc_validated", file=path, misses=len(all_miss))
        return all_miss, patched
    except Exception as e:
        log.error("validate_doc_failed", file=path, error=str(e))
        raise


def build_report(misses):
    lines = ["| # | Левая | Правая |", "|---|-------|--------|"]
    for i, l, r in misses:
        lines.append(f"| {i} | `{l[:40]}` | `{r[:40]}` |")
    return "\n".join(lines)


def highlight_diffs(src: str, misses):
    dst = src.replace(".", "_patched.")
    shutil.copy(src, dst)
    if dst.endswith(".docx"):
        doc = docx.Document(dst)
        for i, (row_idx, *_) in enumerate(misses):
            tbl = doc.tables[0]
            for run in tbl.rows[row_idx - 1].cells[0].paragraphs[0].runs:
                run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
            for run in tbl.rows[row_idx - 1].cells[1].paragraphs[0].runs:
                run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.RED
        doc.save(dst)
    else:
        doc = fitz.open(dst)
        for page in doc:
            page.add_text_annot((50, 50), "⛔️ Есть расхождение")
        doc.save(dst)
    return dst
