from pathlib import Path
import fitz  # PyMuPDF
import docx
import structlog

log = structlog.get_logger(__name__)


def extract_text(path: Path) -> str:
    """Извлекает текст из PDF или DOCX файла"""
    try:
        log.info("extract_text_started", path=str(path), file_size=path.stat().st_size if path.exists() else 0)

        if path.suffix.lower() == ".pdf":
            doc = fitz.open(str(path))
            text = ""
            try:
                # Основной метод - извлечение текста
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    page_text = page.get_text()
                    if page_text.strip():
                        text += page_text + "\n"

                # Если текст пустой, пробуем альтернативный метод
                if not text.strip():
                    log.info("pdf_primary_method_failed", path=str(path))
                    for page_num in range(len(doc)):
                        page = doc.load_page(page_num)
                        # Пробуем извлечь текст как HTML и затем очистить
                        html_text = page.get_text("html")
                        if html_text:
                            # Простая очистка HTML тегов
                            import re

                            clean_text = re.sub(r"<[^>]+>", "", html_text)
                            if clean_text.strip():
                                text += clean_text + "\n"

                log.info("pdf_text_extracted", path=str(path), length=len(text), pages=len(doc))
                return text
            finally:
                doc.close()

        elif path.suffix.lower() in {".docx", ".doc"}:
            d = docx.Document(str(path))

            # Извлекаем текст из параграфов
            paragraphs = []
            for p in d.paragraphs:
                if p.text.strip():
                    paragraphs.append(p.text.strip())

            # Если параграфы пустые, пробуем извлечь из таблиц
            if not paragraphs:
                for table in d.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if paragraph.text.strip():
                                    paragraphs.append(paragraph.text.strip())

            # Если все еще пусто, пробуем извлечь из заголовков
            if not paragraphs:
                for section in d.sections:
                    if section.header.paragraphs:
                        for p in section.header.paragraphs:
                            if p.text.strip():
                                paragraphs.append(p.text.strip())

            text = "\n".join(paragraphs)

            log.info(
                "docx_text_extracted",
                path=str(path),
                length=len(text),
                paragraphs_count=len(paragraphs),
                total_paragraphs=len(d.paragraphs),
                paragraphs=paragraphs[:3],
                all_paragraphs_preview=[p.text[:50] for p in d.paragraphs[:5]],
            )  # Логируем первые 50 символов каждого параграфа
            return text

        else:
            raise ValueError(f"Неподдерживаемый формат файла: {path.suffix}")

    except Exception as e:
        log.error("text_extraction_failed", path=str(path), error=str(e))
        raise ValueError(f"Не удалось извлечь текст из файла: {e}")


def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    """Извлекает текст из байтов файла"""
    import tempfile
    import os

    suffix = Path(filename).suffix.lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(file_bytes)
        tmp_path = Path(tmp.name)

    try:
        text = extract_text(tmp_path)
        return text
    finally:
        if tmp_path.exists():
            os.unlink(tmp_path)
